from __future__ import annotations

import io
import zipfile
from xml.etree import ElementTree
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from autotype.actions import KeyPress, TypeText
from autotype.input import (
    CellBlock,
    DocumentContent,
    InputError,
    ParagraphBlock,
    TableBlock,
    TextRun,
    load_input_content,
    load_input_text,
)


def _write_docx_with_table(path: Path) -> None:
    document = Document()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "Caf\u00e9"
    table.cell(1, 1).text = ""
    document.add_paragraph("Outro")
    document.save(path)


def _write_formatted_docx(path: Path) -> None:
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.add_run("Plain ")
    bold_run = paragraph.add_run("Bold")
    bold_run.bold = True
    italic_run = paragraph.add_run(" Italic")
    italic_run.italic = True
    underline_run = paragraph.add_run(" Underline")
    underline_run.underline = True

    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_run = cell_paragraph.add_run("Cell")
    cell_run.bold = True

    outro = document.add_paragraph()
    outro.add_run("End")

    document.save(path)


def _numbering_num_ids(document: Document, num_fmt: str) -> list[int]:
    root = document.part.numbering_part.numbering_definitions._numbering
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    num_ids: list[int] = []
    for num in root.findall("w:num", ns):
        num_id = int(num.get(qn("w:numId")))
        abstract_id = int(num.find("w:abstractNumId", ns).get(qn("w:val")))
        abstract = root.find(f"w:abstractNum[@w:abstractNumId='{abstract_id}']", ns)
        if abstract is None:
            continue
        level = abstract.find("w:lvl", ns)
        if level is None:
            continue
        num_fmt_element = level.find("w:numFmt", ns)
        if num_fmt_element is None:
            continue
        if num_fmt_element.get(qn("w:val")) == num_fmt:
            num_ids.append(num_id)
    return num_ids


def _set_numbering(paragraph, num_id: int, level: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def _write_order_mixed_docx(path: Path) -> None:
    document = Document()
    decimal_num_id = _numbering_num_ids(document, "decimal")[0]
    bullet_num_id = _numbering_num_ids(document, "bullet")[0]

    document.add_paragraph("Intro")

    bullet = document.add_paragraph("Bullet")
    _set_numbering(bullet, bullet_num_id, 0)

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    numbered = document.add_paragraph("Numbered")
    _set_numbering(numbered, decimal_num_id, 0)

    document.add_paragraph("Outro")
    document.save(path)


def _write_numbering_regression_docx(path: Path) -> None:
    document = Document()
    decimal_num_ids = _numbering_num_ids(document, "decimal")
    assert len(decimal_num_ids) >= 2
    primary_num_id = decimal_num_ids[0]
    restart_num_id = decimal_num_ids[1]

    first = document.add_paragraph("Alpha")
    _set_numbering(first, primary_num_id, 0)

    document.add_paragraph("Interlude")

    second = document.add_paragraph("Bravo")
    _set_numbering(second, primary_num_id, 0)

    nested = document.add_paragraph("Charlie")
    _set_numbering(nested, primary_num_id, 1)

    third = document.add_paragraph("Delta")
    _set_numbering(third, primary_num_id, 0)

    restart = document.add_paragraph("Echo")
    _set_numbering(restart, restart_num_id, 0)

    document.save(path)


def _write_list_formatting_docx(path: Path) -> None:
    document = Document()
    decimal_num_id = _numbering_num_ids(document, "decimal")[0]

    paragraph = document.add_paragraph()
    _set_numbering(paragraph, decimal_num_id, 0)
    paragraph.add_run("Plain ")
    bold_run = paragraph.add_run("Bold")
    bold_run.bold = True
    italic_run = paragraph.add_run(" Italic")
    italic_run.italic = True
    underline_run = paragraph.add_run(" Underline")
    underline_run.underline = True

    document.save(path)


def _write_extended_formatting_docx(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    all_caps = paragraph.add_run("Caps")
    all_caps.font.all_caps = True
    small_caps = paragraph.add_run(" Small")
    small_caps.font.small_caps = True
    superscript = paragraph.add_run(" Super")
    superscript.font.superscript = True
    subscript = paragraph.add_run(" Sub")
    subscript.font.subscript = True

    list_paragraph = document.add_paragraph(style="List Bullet")
    list_run = list_paragraph.add_run("List")
    list_run.font.all_caps = True

    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_run = cell_paragraph.add_run("Cell")
    cell_run.font.subscript = True

    document.save(path)


def _write_conflicting_vertical_alignment_docx(path: Path) -> None:
    document = Document()
    run = document.add_paragraph().add_run("Conflict")
    rPr = run._r.get_or_add_rPr()
    for value in ("superscript", "subscript"):
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), value)
        rPr.append(vert_align)
    document.save(path)


def _write_docx_without_numbering_part(path: Path) -> None:
    source = io.BytesIO()
    Document().add_paragraph("Minimal document").part.package.save(source)
    source.seek(0)

    relationships_path = "word/_rels/document.xml.rels"
    content_types_path = "[Content_Types].xml"
    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    numbering_relationship_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
    with zipfile.ZipFile(source) as archive, zipfile.ZipFile(path, "w") as output:
        for info in archive.infolist():
            if info.filename == "word/numbering.xml":
                continue
            data = archive.read(info.filename)
            if info.filename == relationships_path:
                root = ElementTree.fromstring(data)
                for relationship in root.findall(f"{{{relationship_namespace}}}Relationship"):
                    if relationship.get("Type") == numbering_relationship_type:
                        root.remove(relationship)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            elif info.filename == content_types_path:
                root = ElementTree.fromstring(data)
                for override in root:
                    if override.get("PartName") == "/word/numbering.xml":
                        root.remove(override)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            output.writestr(info, data)


def test_plain_text_normalizes_line_endings() -> None:
    content = load_input_content("Hello\r\nworld", None)

    assert content.blocks == (ParagraphBlock((TextRun("Hello"),)), ParagraphBlock((TextRun("world"),)))
    assert load_input_text("Hello\r\nworld", None) == "Hello\nworld"


def test_txt_file_is_loaded_with_canonical_newlines(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    with path.open("w", encoding="utf-8", newline="") as handle:
        handle.write("alpha\r\nbeta\r\ngamma")

    content = load_input_content(None, path)

    assert content.blocks == (
        ParagraphBlock((TextRun("alpha"),)),
        ParagraphBlock((TextRun("beta"),)),
        ParagraphBlock((TextRun("gamma"),)),
    )
    assert content.to_text() == "alpha\nbeta\ngamma"


def test_docx_paragraphs_and_tables_preserve_document_order(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    _write_docx_with_table(path)

    content = load_input_content(None, path)

    assert isinstance(content.blocks[0], ParagraphBlock)
    assert content.blocks[0].text == "Intro"
    assert isinstance(content.blocks[1], TableBlock)
    assert content.blocks[1].rows[0][0].text == "A"
    assert content.blocks[1].rows[0][1].text == "B"
    assert content.blocks[1].rows[1][0].text == "Caf\u00e9"
    assert content.blocks[1].rows[1][1].text == ""
    assert isinstance(content.blocks[2], ParagraphBlock)
    assert content.blocks[2].text == "Outro"
    assert content.to_text() == "Intro\nA\tB\nCaf\u00e9\t\nOutro"


def test_docx_run_formatting_is_preserved_in_actions(tmp_path: Path) -> None:
    path = tmp_path / "formatted.docx"
    _write_formatted_docx(path)

    content = load_input_content(None, path)
    actions = content.to_actions()

    assert actions == [
        TypeText("Plain "),
        KeyPress("CTRL+B"),
        TypeText("Bold"),
        KeyPress("CTRL+B"),
        KeyPress("CTRL+I"),
        TypeText(" Italic"),
        KeyPress("CTRL+I"),
        KeyPress("CTRL+U"),
        TypeText(" Underline"),
        KeyPress("CTRL+U"),
        TypeText("\n"),
        KeyPress("CTRL+B"),
        TypeText("Cell"),
        KeyPress("CTRL+B"),
        TypeText("\n"),
        TypeText("End"),
    ]


def test_docx_lists_are_linearized_with_order_and_indentation(tmp_path: Path) -> None:
    path = tmp_path / "lists.docx"
    _write_order_mixed_docx(path)

    content = load_input_content(None, path)

    assert isinstance(content.blocks[0], ParagraphBlock)
    assert content.blocks[0].text == "Intro"
    assert isinstance(content.blocks[1], ParagraphBlock)
    assert content.blocks[1].prefix == "\uf0b7 "
    assert content.blocks[1].text == "\uf0b7 Bullet"
    assert isinstance(content.blocks[2], TableBlock)
    assert content.blocks[2].rows[0][0].text == "A"
    assert content.blocks[2].rows[0][1].text == "B"
    assert isinstance(content.blocks[3], ParagraphBlock)
    assert content.blocks[3].prefix == "1. "
    assert content.blocks[3].text == "1. Numbered"
    assert isinstance(content.blocks[4], ParagraphBlock)
    assert content.blocks[4].text == "Outro"
    assert content.to_text() == "Intro\n\uf0b7 Bullet\nA\tB\n1. Numbered\nOutro"


def test_docx_numbered_lists_continue_and_restart_by_numid_and_ilvl(tmp_path: Path) -> None:
    path = tmp_path / "numbering.docx"
    _write_numbering_regression_docx(path)

    content = load_input_content(None, path)

    assert content.to_text() == "1. Alpha\nInterlude\n2. Bravo\n  1. Charlie\n3. Delta\n1. Echo"
    assert content.blocks[0].text == "1. Alpha"
    assert content.blocks[1].text == "Interlude"
    assert content.blocks[2].text == "2. Bravo"
    assert content.blocks[3].text == "  1. Charlie"
    assert content.blocks[4].text == "3. Delta"
    assert content.blocks[5].text == "1. Echo"


def test_docx_list_items_preserve_inline_formatting_in_actions(tmp_path: Path) -> None:
    path = tmp_path / "list-formatting.docx"
    _write_list_formatting_docx(path)

    content = load_input_content(None, path)
    actions = content.to_actions()

    assert actions == [
        TypeText("1. "),
        TypeText("Plain "),
        KeyPress("CTRL+B"),
        TypeText("Bold"),
        KeyPress("CTRL+B"),
        KeyPress("CTRL+I"),
        TypeText(" Italic"),
        KeyPress("CTRL+I"),
        KeyPress("CTRL+U"),
        TypeText(" Underline"),
        KeyPress("CTRL+U"),
    ]


def test_docx_extended_direct_formatting_is_extracted_and_emitted(tmp_path: Path) -> None:
    path = tmp_path / "extended-formatting.docx"
    _write_extended_formatting_docx(path)

    content = load_input_content(None, path)
    paragraph = content.blocks[0]
    assert isinstance(paragraph, ParagraphBlock)
    assert paragraph.runs == (
        TextRun("Caps", all_caps=True),
        TextRun(" Small", small_caps=True),
        TextRun(" Super", superscript=True),
        TextRun(" Sub", subscript=True),
    )
    assert content.has_formatting is True
    assert content.to_actions()[:12] == [
        KeyPress("CTRL+SHIFT+A"),
        TypeText("Caps"),
        KeyPress("CTRL+SHIFT+A"),
        KeyPress("CTRL+SHIFT+K"),
        TypeText(" Small"),
        KeyPress("CTRL+SHIFT+K"),
        KeyPress("CTRL+SHIFT+EQUALS"),
        TypeText(" Super"),
        KeyPress("CTRL+SHIFT+EQUALS"),
        KeyPress("CTRL+EQUALS"),
        TypeText(" Sub"),
        KeyPress("CTRL+EQUALS"),
    ]


def test_docx_extended_formatting_works_in_lists_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "extended-block-formatting.docx"
    _write_extended_formatting_docx(path)

    content = load_input_content(None, path)
    actions = content.to_actions()

    assert KeyPress("CTRL+SHIFT+A") in actions
    assert KeyPress("CTRL+EQUALS") in actions
    assert content.blocks[1].text == "\uf0b7 List"
    assert content.blocks[2].rows[0][0].text == "Cell"


def test_docx_rejects_conflicting_vertical_alignment(tmp_path: Path) -> None:
    path = tmp_path / "conflicting-formatting.docx"
    _write_conflicting_vertical_alignment_docx(path)

    with pytest.raises(InputError, match="both superscript and subscript"):
        load_input_content(None, path)


def test_minimal_docx_without_numbering_part_loads_text(tmp_path: Path) -> None:
    path = tmp_path / "minimal-without-numbering.docx"
    _write_docx_without_numbering_part(path)

    assert load_input_text(None, path) == "Minimal document"


def test_docx_formatting_resets_at_block_boundaries(tmp_path: Path) -> None:
    path = tmp_path / "boundary-formatting.docx"
    document = Document()
    first = document.add_paragraph()
    first_run = first.add_run("First")
    first_run.font.superscript = True
    second = document.add_paragraph()
    second_run = second.add_run("Second")
    second_run.font.all_caps = True
    document.save(path)

    content = load_input_content(None, path)

    assert content.to_actions() == [
        KeyPress("CTRL+SHIFT+EQUALS"),
        TypeText("First"),
        KeyPress("CTRL+SHIFT+EQUALS"),
        TypeText("\n"),
        KeyPress("CTRL+SHIFT+A"),
        TypeText("Second"),
        KeyPress("CTRL+SHIFT+A"),
    ]


@pytest.mark.parametrize("field_name", ["all_caps", "small_caps", "superscript", "subscript"])
def test_extended_formatting_counts_as_formatting(field_name: str) -> None:
    run = TextRun("x", **{field_name: True})
    content = DocumentContent((ParagraphBlock((run,)),), source_kind="docx")

    assert content.has_formatting is True


def test_docx_empty_document_normalizes_to_empty_text(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    assert load_input_text(None, path) == ""


def test_unsupported_file_extension_raises_input_error(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    path.write_text("irrelevant", encoding="utf-8")

    with pytest.raises(InputError, match="Unsupported input file type"):
        load_input_content(None, path)


def test_malformed_docx_raises_input_error(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not-a-valid-docx")

    with pytest.raises(InputError, match="Could not read DOCX file"):
        load_input_text(None, path)


def test_missing_explicit_file_raises_input_error(tmp_path: Path) -> None:
    path = tmp_path / "missing.txt"

    with pytest.raises(InputError, match="does not exist"):
        load_input_text(None, path)
