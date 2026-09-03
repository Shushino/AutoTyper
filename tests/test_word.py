from __future__ import annotations

import importlib.util
import os
from pathlib import Path

import pytest
from docx import Document

import autotype.cli as cli_module
import autotype.word as word_module
from autotype.cli import build_parser, main
from autotype.word import WordDocumentInserter, WordPreflightError


class FakeInformation:
    def __init__(self, inside_table: bool = False) -> None:
        self.inside_table = inside_table

    def __call__(self, index: int) -> bool:
        assert index == 12
        return self.inside_table


class FakeSelection:
    def __init__(self, *, start: int = 1, end: int = 1, story_type: int = 1, inside_table: bool = False) -> None:
        self.Start = start
        self.End = end
        self.StoryType = story_type
        self.Information = FakeInformation(inside_table)
        self.inserted: tuple[str, dict[str, object]] | None = None

    def InsertFile(self, path: str, **kwargs: object) -> None:
        self.inserted = (path, kwargs)


class FakeDocument:
    def __init__(self, *, read_only: bool = False, protection_type: int = -1) -> None:
        self.ReadOnly = read_only
        self.ProtectionType = protection_type


class FakeApplication:
    def __init__(self, document: FakeDocument | None, selection: FakeSelection | None) -> None:
        self.ActiveDocument = document
        self.Selection = selection


def _source_docx(tmp_path: Path) -> Path:
    path = tmp_path / "source.docx"
    Document().add_paragraph("source").part.package.save(path)
    return path


def _inserter(app: FakeApplication) -> WordDocumentInserter:
    return WordDocumentInserter(active_object=lambda name: app)


def test_keyboard_target_remains_default() -> None:
    assert build_parser().parse_args([]).target == "keyboard"


def test_word_preflight_rejects_unavailable_word(tmp_path: Path) -> None:
    source = _source_docx(tmp_path)

    def unavailable(_: str):
        raise OSError("not running")

    with pytest.raises(WordPreflightError, match="will not launch"):
        WordDocumentInserter(active_object=unavailable).preflight(source)


@pytest.mark.parametrize(
    ("document", "selection", "message"),
    [
        (None, FakeSelection(), "no active document"),
        (FakeDocument(read_only=True), FakeSelection(), "read-only"),
        (FakeDocument(protection_type=0), FakeSelection(), "protected"),
        (FakeDocument(), FakeSelection(start=1, end=2), "collapsed"),
        (FakeDocument(), FakeSelection(story_type=6), "main document body"),
        (FakeDocument(), FakeSelection(inside_table=True), "inside a Word table"),
    ],
)
def test_word_preflight_rejects_unsafe_context(
    tmp_path: Path,
    document: FakeDocument | None,
    selection: FakeSelection,
    message: str,
) -> None:
    source = _source_docx(tmp_path)
    app = FakeApplication(document, selection)

    with pytest.raises(WordPreflightError, match=message):
        _inserter(app).preflight(source)


def test_word_insertion_uses_exact_native_insertfile_arguments(tmp_path: Path) -> None:
    source = _source_docx(tmp_path)
    selection = FakeSelection()
    app = FakeApplication(FakeDocument(), selection)

    _inserter(app).insert(source)

    assert selection.inserted == (
        str(source.resolve()),
        {"ConfirmConversions": False, "Link": False, "Attachment": False},
    )


def test_word_com_failure_explains_undo(tmp_path: Path) -> None:
    source = _source_docx(tmp_path)
    selection = FakeSelection()

    def fail_insert(_: str, **__: object) -> None:
        raise OSError("COM failed")

    selection.InsertFile = fail_insert  # type: ignore[method-assign]
    app = FakeApplication(FakeDocument(), selection)

    with pytest.raises(WordPreflightError, match="partially inserted.*Undo"):
        _inserter(app).insert(source)


def test_word_mode_rejects_non_docx_and_plain_text(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))
    txt = tmp_path / "source.txt"
    txt.write_text("text", encoding="utf-8")

    with pytest.raises(SystemExit, match=r"\.docx"):
        main(["--target", "word", "plain text"])
    with pytest.raises(SystemExit, match=r"\.docx"):
        main(["--target", "word", str(txt)])


def test_word_mode_rejects_malformed_docx_even_in_dry_run(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))
    malformed = tmp_path / "malformed.docx"
    malformed.write_bytes(b"not a zip package")

    with pytest.raises(SystemExit, match="valid DOCX"):
        main(["--target", "word", "--dry-run", str(malformed)])


def test_word_dry_run_does_not_construct_or_contact_adapter(tmp_path: Path, monkeypatch: pytest.MonkeyPatch, capsys) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))
    source = _source_docx(tmp_path)

    class ExplodingInserter:
        def __init__(self) -> None:
            raise AssertionError("Word adapter must not be constructed during dry-run")

    monkeypatch.setattr(cli_module, "WordDocumentInserter", ExplodingInserter)
    assert main(["--target", "word", "--dry-run", str(source)]) == 0
    output = capsys.readouterr().out
    assert "WORD FIDELITY DRY RUN" in output
    assert "not imported or contacted" in output
    assert "No Word document was changed" in output


def test_keyboard_dry_run_still_uses_existing_pipeline(tmp_path: Path, monkeypatch: pytest.MonkeyPatch, capsys) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))
    source = _source_docx(tmp_path)

    assert main(["--target", "keyboard", "--dry-run", "--countdown", "0", str(source)]) == 0
    assert "[DRY RUN]" in capsys.readouterr().out


@pytest.mark.word_integration
def test_word_integration_smoke_fixture() -> None:
    if os.name != "nt":
        pytest.skip("Microsoft Word integration requires Windows")
    if importlib.util.find_spec("win32com") is None:
        pytest.skip("pywin32 is not installed")
    if os.environ.get("AUTOTYPER_WORD_INTEGRATION") != "1":
        pytest.skip("Set AUTOTYPER_WORD_INTEGRATION=1 to opt in")

    fixture = Path(
        os.environ.get(
            "AUTOTYPER_WORD_FIXTURE",
            str(Path.home() / "Desktop" / "files" / "AutoTyper_Full_Capability_Test.docx"),
        )
    ).expanduser()
    if not fixture.is_file():
        pytest.skip(f"Word fixture not found: {fixture}")

    from win32com.client import GetActiveObject

    try:
        application = GetActiveObject("Word.Application")
    except Exception:
        pytest.skip("Microsoft Word is not already running")

    target = application.Documents.Add()
    try:
        target.Activate()
        target.Range(0, 0).Select()
        WordDocumentInserter().insert(fixture)

        assert int(target.Tables.Count) == 4
        assert int(target.Tables.Item(3).Rows.Item(1).Cells.Count) == 2
        assert any(int(paragraph.Range.ListFormat.ListType) != 0 for paragraph in target.Paragraphs)

        headings = {
            str(paragraph.Range.Text).rstrip("\r\a"): str(paragraph.Style.NameLocal)
            for paragraph in target.Paragraphs
        }
        assert headings["AutoTyper Full Capability Smoke Test"] == "Heading 1"
        assert headings["1. Plain paragraphs"] == "Heading 2"

        for paragraph in target.Paragraphs:
            if str(paragraph.Range.Text).startswith("Paragraph after the list."):
                assert int(paragraph.Range.ListFormat.ListType) == 0
                break
        else:
            pytest.fail("fixture paragraph after list was not found")
    finally:
        target.Close(SaveChanges=0)
