from __future__ import annotations

import json
from pathlib import Path

import autotype.cli as cli_module
from docx import Document
import pytest

from autotype.cli import build_parser, main
from autotype.executors import MockExecutor


def _write_table_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    document.add_paragraph("Outro")
    document.save(path)


def _write_merged_table_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Merged"
    table.cell(0, 0).merge(table.cell(0, 1))
    document.save(path)


def _write_formatted_docx(path: Path) -> None:
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.add_run("Plain ")
    bold_run = paragraph.add_run("Bold")
    bold_run.bold = True
    paragraph.add_run(" ")
    italic_run = paragraph.add_run("Italic")
    italic_run.italic = True
    paragraph.add_run(" ")
    underline_run = paragraph.add_run("Underline")
    underline_run.underline = True

    table = document.add_table(rows=1, cols=1)
    table_run = table.cell(0, 0).paragraphs[0].add_run("Cell")
    table_run.bold = True

    document.save(path)


def _write_list_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Intro")
    document.add_paragraph("Bullet", style="List Bullet")
    document.add_paragraph("Numbered", style="List Number")
    document.save(path)


def test_parser_accepts_text_and_controls() -> None:
    parser = build_parser()
    args = parser.parse_args(
        [
            "hello",
            "--speed",
            "55",
            "--countdown",
            "2",
            "--profile",
            "careful",
            "--seed",
            "1234",
            "--typo-rate",
            "0.05",
            "--dry-run",
            "--pause-key",
            "F7",
            "--stop-key",
            "F11",
        ]
    )

    assert args.text == "hello"
    assert args.speed == 55.0
    assert args.countdown == 2.0
    assert args.profile == "careful"
    assert args.seed == 1234
    assert args.typo_rate == 0.05
    assert args.dry_run is True
    assert not hasattr(args, "progress")
    assert args.pause_key == "F7"
    assert args.stop_key == "F11"


def test_parser_help_includes_examples_and_progress_flag() -> None:
    help_text = build_parser().format_help()

    assert "Examples:" in help_text
    assert "autotype \"Hello world\" --dry-run --seed 1234" in help_text
    assert "autotype --file lists.docx --dry-run --profile natural --seed 1234" in help_text
    assert "autotype --file lists.docx --countdown 5 --progress" in help_text
    assert "autotype --show-config" in help_text
    assert "autotype --speed 110 --save-config --config .pytest-tmp/custom_config.json" in help_text
    assert "--config" in help_text
    assert "--show-config" in help_text
    assert "--save-config" in help_text
    assert "--no-progress" in help_text
    assert "--progress" in help_text


def test_dry_run_cli_returns_success(monkeypatch, tmp_path: Path, capsys) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))

    exit_code = main(
        [
            "hello",
            "--dry-run",
            "--speed",
            "60",
            "--countdown",
            "0",
            "--profile",
            "precise",
            "--seed",
            "7",
            "--typo-rate",
            "0.0",
        ]
    )

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "[DRY RUN]" in captured.out
    assert "Profile: precise" in captured.out
    assert "Typo rate: 0.000" in captured.out
    assert "Seed: 7" in captured.out
    assert "TypeText('h')" in captured.out
    assert "Pause(" in captured.out
    assert "Estimated total duration:" in captured.out
    assert "Summary:" in captured.out
    assert "Input kind: plain text" in captured.out


def test_show_config_prints_effective_defaults_and_skips_execution(monkeypatch, tmp_path: Path, capsys) -> None:
    monkeypatch.setenv("APPDATA", str(tmp_path))
    monkeypatch.setattr(cli_module, "read_input_content", lambda args: pytest.fail("show-config should not read input"))

    exit_code = main(["--show-config"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert json.loads(captured.out) == {
        "version": 1,
        "profile": "natural",
        "speed": 40.0,
        "typo_rate": 0.0,
        "countdown": 5.0,
        "progress": False,
    }


def test_config_values_override_built_in_defaults_and_cli_overrides_config(tmp_path: Path, capsys) -> None:
    path = tmp_path / "config.json"
    path.write_text(
        json.dumps(
            {
                "version": 1,
                "profile": "careful",
                "speed": 75,
                "typo_rate": 0.04,
                "countdown": 2,
                "progress": True,
            }
        ),
        encoding="utf-8",
    )

    exit_code = main([
        "--config",
        str(path),
        "--speed",
        "90",
        "--no-progress",
        "--show-config",
    ])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert json.loads(captured.out) == {
        "version": 1,
        "profile": "careful",
        "speed": 90.0,
        "typo_rate": 0.04,
        "countdown": 2.0,
        "progress": False,
    }


def test_save_config_persists_effective_settings(tmp_path: Path, capsys) -> None:
    path = tmp_path / "saved.json"

    exit_code = main([
        "--speed",
        "110",
        "--save-config",
        "--config",
        str(path),
    ])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Saved configuration to" in captured.out
    assert json.loads(path.read_text(encoding="utf-8")) == {
        "version": 1,
        "profile": "natural",
        "speed": 110.0,
        "typo_rate": 0.0,
        "countdown": 5.0,
        "progress": False,
    }


def test_missing_explicit_config_fails(tmp_path: Path) -> None:
    missing = tmp_path / "missing.json"

    with pytest.raises(SystemExit, match="Configuration file does not exist"):
        main(["--config", str(missing), "--show-config"])


def test_dry_run_cli_reports_docx_summary_and_formatting(tmp_path: Path, capsys) -> None:
    path = tmp_path / "formatted.docx"
    _write_formatted_docx(path)

    exit_code = main(["--file", str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise", "--seed", "7"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Input kind: DOCX" in captured.out
    assert "Formatting toggles: 8" in captured.out
    assert "Characters: 32" in captured.out
    assert "KeyPress('CTRL+B')" in captured.out
    assert "KeyPress('CTRL+I')" in captured.out
    assert "KeyPress('CTRL+U')" in captured.out


def test_dry_run_cli_reports_docx_list_prefixes(tmp_path: Path, capsys) -> None:
    path = tmp_path / "lists.docx"
    _write_list_docx(path)

    exit_code = main(["--file", str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise", "--seed", "7"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Input kind: DOCX" in captured.out
    assert "TypeText('\u2022')" in captured.out
    assert "TypeText('\\uf0b7')" not in captured.out
    assert "TypeText('1')" in captured.out
    assert "TypeText('.')" in captured.out


def test_dry_run_cli_does_not_duplicate_merged_cell_text(tmp_path: Path, capsys) -> None:
    path = tmp_path / "merged-table.docx"
    _write_merged_table_docx(path)

    exit_code = main(["--file", str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert captured.out.count("TypeText('M')") == 1


def test_dry_run_cli_reports_extended_formatting(tmp_path: Path, capsys) -> None:
    path = tmp_path / "extended-formatting.docx"
    document = Document()
    paragraph = document.add_paragraph()
    for text, attribute in (
        ("Caps", "all_caps"),
        ("Small", "small_caps"),
        ("Super", "superscript"),
        ("Sub", "subscript"),
    ):
        run = paragraph.add_run(text)
        setattr(run.font, attribute, True)
    document.save(path)

    exit_code = main(["--file", str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise", "--seed", "7"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Formatting toggles: 8" in captured.out
    assert "KeyPress('CTRL+SHIFT+A')" in captured.out
    assert "KeyPress('CTRL+SHIFT+K')" in captured.out
    assert "KeyPress('CTRL+SHIFT+EQUALS')" in captured.out
    assert "KeyPress('CTRL+SHIFT+MINUS')" in captured.out


def test_dry_run_cli_accepts_table_docx_positional_input(tmp_path: Path, capsys) -> None:
    path = tmp_path / "sample.docx"
    _write_table_docx(path)

    exit_code = main([str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise", "--seed", "7"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "TypeText('I')" in captured.out
    assert "TypeText('\\t')" in captured.out
    assert "TypeText('O')" in captured.out


def test_dry_run_cli_accepts_table_docx_via_file_flag(tmp_path: Path, capsys) -> None:
    path = tmp_path / "sample.docx"
    _write_table_docx(path)

    exit_code = main(["--file", str(path), "--dry-run", "--speed", "60", "--countdown", "0", "--profile", "precise", "--seed", "7"])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "TypeText('\\t')" in captured.out
    assert "TypeText('D')" in captured.out


class _DummyHotkeyMonitor:
    def __init__(self, *args, **kwargs) -> None:
        pass

    def start(self) -> None:
        return None

    def close(self) -> None:
        return None


def test_live_run_prints_friendly_status_and_progress(monkeypatch, capsys) -> None:
    monkeypatch.setattr(cli_module, "WindowsExecutor", MockExecutor)
    monkeypatch.setattr(cli_module, "WindowsHotkeyMonitor", _DummyHotkeyMonitor)

    exit_code = main([
        "hello",
        "--countdown",
        "0",
        "--profile",
        "precise",
        "--seed",
        "7",
        "--progress",
    ])

    captured = capsys.readouterr()
    assert exit_code == 0
    assert "Running" in captured.out
    assert "Finished" in captured.out
    assert "Progress:" in captured.out
    assert "Idle" not in captured.out
