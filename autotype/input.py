from __future__ import annotations

import re
from io import BytesIO
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .actions import Action, KeyPress, TypeText


SUPPORTED_INPUT_SUFFIXES = {".txt", ".docx"}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


class InputError(ValueError):
    """Raised when user input cannot be resolved into canonical text."""


@dataclass(frozen=True, slots=True)
class TextRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    all_caps: bool = False
    small_caps: bool = False
    superscript: bool = False
    subscript: bool = False


@dataclass(frozen=True, slots=True)
class ParagraphBlock:
    runs: tuple[TextRun, ...]
    prefix: str = ""

    @property
    def text(self) -> str:
        return f"{self.prefix}{''.join(run.text for run in self.runs)}"


@dataclass(frozen=True, slots=True)
class CellBlock:
    paragraphs: tuple[ParagraphBlock, ...]

    @property
    def text(self) -> str:
        return "\n".join(paragraph.text for paragraph in self.paragraphs)


@dataclass(frozen=True, slots=True)
class TableBlock:
    rows: tuple[tuple[CellBlock, ...], ...]


DocumentBlock = ParagraphBlock | TableBlock


@dataclass(frozen=True, slots=True)
class DocumentContent:
    """Normalized textual content extracted from a document or text file."""

    blocks: tuple[DocumentBlock, ...]
    source_kind: str = "text"

    @classmethod
    def from_text(cls, text: str) -> "DocumentContent":
        normalized = _normalize_line_endings(text)
        return cls(tuple(ParagraphBlock((TextRun(line),)) for line in normalized.split("\n")))

    @classmethod
    def from_paragraphs(cls, paragraphs: Iterable[str]) -> "DocumentContent":
        return cls(tuple(ParagraphBlock((TextRun(_normalize_line_endings(paragraph)),)) for paragraph in paragraphs))

    @classmethod
    def from_docx(cls, document: DocxDocument) -> "DocumentContent":
        return cls(tuple(_iter_docx_blocks(document)), source_kind="docx")

    @property
    def paragraphs(self) -> tuple[str, ...]:
        return tuple(block.text for block in self.blocks if isinstance(block, ParagraphBlock))

    @property
    def has_formatting(self) -> bool:
        if self.source_kind != "docx":
            return False
        for block in self.blocks:
            if isinstance(block, ParagraphBlock):
                if any(_run_has_formatting(run) for run in block.runs):
                    return True
            elif isinstance(block, TableBlock):
                for row in block.rows:
                    for cell in row:
                        for paragraph in cell.paragraphs:
                            if any(_run_has_formatting(run) for run in paragraph.runs):
                                return True
        return False

    def to_text(self) -> str:
        parts: list[str] = []
        for block in self.blocks:
            if isinstance(block, ParagraphBlock):
                parts.append(block.text)
            else:
                parts.append("\n".join("\t".join(cell.text for cell in row) for row in block.rows))
        return "\n".join(parts)

    def to_actions(self) -> list[Action]:
        actions: list[Action] = []
        current_style = _StyleState()

        for index, block in enumerate(self.blocks):
            if isinstance(block, ParagraphBlock):
                current_style = _emit_paragraph_actions(actions, block, current_style)
            else:
                current_style = _emit_table_actions(actions, block)
            if index + 1 < len(self.blocks):
                current_style = _emit_style_transition(actions, current_style, _StyleState())
                actions.append(TypeText("\n"))

        if current_style != _StyleState():
            actions.extend(_style_close_actions(current_style))

        return actions


def load_input_text(raw_text: str | None, file_path: Path | None) -> str:
    return load_input_content(raw_text, file_path).to_text()


def load_input_content(raw_text: str | None, file_path: Path | None) -> DocumentContent:
    if file_path is not None:
        return load_content_from_path(file_path)
    if raw_text is None:
        raise InputError("Provide either text or --file.")
    if raw_text == "":
        return DocumentContent.from_text("")

    candidate = Path(raw_text).expanduser()
    if candidate.exists():
        return load_content_from_path(candidate)

    return DocumentContent.from_text(raw_text)


def load_content_from_path(path: Path) -> DocumentContent:
    resolved = path.expanduser()
    if not resolved.exists():
        raise InputError(f"Input file does not exist: {resolved}")
    if not resolved.is_file():
        raise InputError(f"Input path is not a file: {resolved}")

    suffix = resolved.suffix.lower()
    if suffix == ".txt":
        return _load_txt_content(resolved)
    if suffix == ".docx":
        return _load_docx_content(resolved)

    supported = ", ".join(sorted(SUPPORTED_INPUT_SUFFIXES))
    raise InputError(f"Unsupported input file type: {resolved.suffix or '<none>'}. Supported types: {supported}")


def _load_txt_content(path: Path) -> DocumentContent:
    try:
        text = path.read_text(encoding="utf-8-sig")
    except OSError as exc:  # pragma: no cover - exercised via higher-level tests
        raise InputError(f"Could not read text file: {path}") from exc
    return DocumentContent.from_text(text)


def _load_docx_content(path: Path) -> DocumentContent:
    try:
        document = Document(BytesIO(path.read_bytes()))
    except Exception as exc:  # pragma: no cover - exercised via malformed DOCX tests
        raise InputError(f"Could not read DOCX file: {path}") from exc

    return DocumentContent.from_docx(document)


def _iter_docx_blocks(document: DocxDocument) -> Iterator[DocumentBlock]:
    list_resolver = _DocxListResolver(document)
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            yield ParagraphBlock(tuple(_iter_paragraph_runs(paragraph)), prefix=list_resolver.prefix_for(paragraph))
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            yield TableBlock(
                tuple(
                    tuple(_iter_cell_block(cell, list_resolver) for cell in row.cells)
                    for row in table.rows
                )
            )


def _iter_paragraph_runs(paragraph: Paragraph) -> Iterator[TextRun]:
    for run in paragraph.runs:
        font = run.font
        superscript, subscript = _direct_vertical_alignment(run)
        yield TextRun(
            text=_normalize_line_endings(run.text or ""),
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
            all_caps=font.all_caps is True,
            small_caps=font.small_caps is True,
            superscript=superscript,
            subscript=subscript,
        )


def _direct_vertical_alignment(run) -> tuple[bool, bool]:
    rPr = run._r.rPr
    values = set()
    if rPr is not None:
        values = {
            element.get(qn("w:val"))
            for element in rPr.findall("w:vertAlign", NS)
        }

    superscript = "superscript" in values
    subscript = "subscript" in values
    if superscript and subscript:
        raise InputError("DOCX run cannot be both superscript and subscript")
    return superscript, subscript


def _iter_cell_block(cell, list_resolver: "_DocxListResolver") -> CellBlock:
    return CellBlock(
        tuple(
            ParagraphBlock(tuple(_iter_paragraph_runs(paragraph)), prefix=list_resolver.prefix_for(paragraph))
            for paragraph in cell.paragraphs
        )
    )


@dataclass(frozen=True, slots=True)
class _ListLevelDefinition:
    num_fmt: str
    lvl_text: str
    start: int = 1


@dataclass(frozen=True, slots=True)
class _ListTemplate:
    level_definitions: dict[int, _ListLevelDefinition]


@dataclass(frozen=True, slots=True)
class _ListStyleTemplate:
    template: _ListTemplate
    level: int
    definition: _ListLevelDefinition


@dataclass(slots=True)
class _ListState:
    counters: dict[int, int] = field(default_factory=dict)


class _DocxListResolver:
    def __init__(self, document: DocxDocument) -> None:
        self._templates_by_num_id, self._templates_by_style_id = _build_list_templates(document)
        self._states: dict[tuple[str, str | int], _ListState] = {}

    def prefix_for(self, paragraph: Paragraph) -> str:
        spec = self._resolve_spec(paragraph)
        if spec is None:
            return ""

        key, template, level, definition = spec
        state = self._states.setdefault(key, _ListState())
        state.counters[level] = state.counters.get(level, definition.start - 1) + 1
        for deeper_level in list(state.counters):
            if deeper_level > level:
                del state.counters[deeper_level]

        label = _render_list_label(template, level, state.counters)
        indent = "  " * level
        return f"{indent}{label} "

    def _resolve_spec(
        self, paragraph: Paragraph
    ) -> tuple[tuple[str, str | int], _ListTemplate, int, _ListLevelDefinition] | None:
        explicit = _explicit_list_spec(paragraph, self._templates_by_num_id)
        if explicit is not None:
            return explicit

        style = getattr(paragraph, "style", None)
        if style is None:
            return None

        style_id = getattr(style, "style_id", None)
        if not style_id:
            return None

        template = self._templates_by_style_id.get(style_id)
        if template is not None:
            return (("style", style_id), template.template, template.level, template.definition)

        heuristic = _heuristic_style_spec(style_id, getattr(style, "name", ""))
        if heuristic is None:
            return None

        level, definition = heuristic
        template = _ListTemplate({level: definition})
        return (("style", style_id), template, level, definition)


def _build_list_templates(document: DocxDocument) -> tuple[dict[int, _ListTemplate], dict[str, _ListStyleTemplate]]:
    numbering_relationship = next(
        (relationship for relationship in document.part.rels.values() if relationship.reltype == RT.NUMBERING),
        None,
    )
    if numbering_relationship is None:
        return {}, {}

    root = numbering_relationship.target_part.numbering_definitions._numbering
    abstract_templates: dict[int, _ListTemplate] = {}
    style_templates: dict[str, _ListStyleTemplate] = {}
    num_templates: dict[int, _ListTemplate] = {}

    for abstract in root.findall("w:abstractNum", NS):
        abstract_id = int(abstract.get(qn("w:abstractNumId")))
        level_definitions: dict[int, _ListLevelDefinition] = {}

        for level in abstract.findall("w:lvl", NS):
            ilvl = int(level.get(qn("w:ilvl")))
            num_fmt = _xml_attr(level.find("w:numFmt", NS), qn("w:val"), default="decimal")
            lvl_text = _xml_attr(level.find("w:lvlText", NS), qn("w:val"), default="%1.")
            start = int(_xml_attr(level.find("w:start", NS), qn("w:val"), default="1"))
            definition = _ListLevelDefinition(num_fmt=num_fmt, lvl_text=lvl_text, start=start)
            level_definitions[ilvl] = definition

            style_id = _xml_attr(level.find("w:pStyle", NS), qn("w:val"), default="")
            if style_id:
                style_templates[style_id] = _ListStyleTemplate(_ListTemplate(level_definitions), ilvl, definition)

        abstract_templates[abstract_id] = _ListTemplate(level_definitions)

    for num in root.findall("w:num", NS):
        num_id = int(num.get(qn("w:numId")))
        abstract_id = int(_xml_attr(num.find("w:abstractNumId", NS), qn("w:val"), default="0"))
        template = abstract_templates.get(abstract_id)
        if template is not None:
            num_templates[num_id] = template

    return num_templates, style_templates


def _explicit_list_spec(
    paragraph: Paragraph,
    templates_by_num_id: dict[int, _ListTemplate],
) -> tuple[tuple[str, str | int], _ListTemplate, int, _ListLevelDefinition] | None:
    pPr = paragraph._p.pPr
    if pPr is None or pPr.numPr is None:
        return None

    num_id = pPr.numPr.numId.val if pPr.numPr.numId is not None else None
    ilvl = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else None
    if num_id is None or ilvl is None:
        return None

    template = templates_by_num_id.get(int(num_id))
    if template is None:
        return None

    level = int(ilvl)
    definition = _resolve_level_definition(template, level)
    return (("num", int(num_id)), template, level, definition)


def _resolve_level_definition(template: _ListTemplate, level: int) -> _ListLevelDefinition:
    if level in template.level_definitions:
        return template.level_definitions[level]

    if not template.level_definitions:
        return _ListLevelDefinition(num_fmt="decimal", lvl_text=f"%{level + 1}.")

    lower_levels = [candidate for candidate in template.level_definitions if candidate <= level]
    if lower_levels:
        base = template.level_definitions[max(lower_levels)]
    else:
        base = template.level_definitions[min(template.level_definitions)]

    if base.num_fmt == "bullet":
        return _ListLevelDefinition(num_fmt="bullet", lvl_text=base.lvl_text or "•", start=base.start)
    return _ListLevelDefinition(num_fmt=base.num_fmt, lvl_text=f"%{level + 1}.", start=base.start)


def _heuristic_style_spec(style_id: str, style_name: str) -> tuple[int, _ListLevelDefinition] | None:
    normalized = style_name.lower().strip()
    if normalized.startswith("list bullet"):
        level = _style_name_level(style_name)
        return level, _ListLevelDefinition(num_fmt="bullet", lvl_text="•")
    if normalized.startswith("list number"):
        level = _style_name_level(style_name)
        return level, _ListLevelDefinition(num_fmt="decimal", lvl_text="%1.")
    return None


def _style_name_level(style_name: str) -> int:
    match = re.search(r"\b(\d+)\b", style_name)
    if match is None:
        return 0
    return max(0, int(match.group(1)) - 1)


def _render_list_label(template: _ListTemplate, level: int, counters: dict[int, int]) -> str:
    definition = _resolve_level_definition(template, level)
    if definition.num_fmt == "bullet":
        return definition.lvl_text or "•"

    def replace_placeholder(match: re.Match[str]) -> str:
        placeholder_level = int(match.group(1)) - 1
        placeholder_definition = _resolve_level_definition(template, placeholder_level)
        value = counters.get(placeholder_level, placeholder_definition.start)
        return _format_list_value(value, placeholder_definition.num_fmt)

    label = re.sub(r"%([1-9]\d*)", replace_placeholder, definition.lvl_text)
    if "%" not in definition.lvl_text:
        return label
    return label


def _format_list_value(value: int, num_fmt: str) -> str:
    if num_fmt == "decimal":
        return str(value)
    if num_fmt == "bullet":
        return "•"
    if num_fmt == "lowerLetter":
        return _alpha_sequence(value, lowercase=True)
    if num_fmt == "upperLetter":
        return _alpha_sequence(value, lowercase=False)
    if num_fmt == "lowerRoman":
        return _roman_sequence(value).lower()
    if num_fmt == "upperRoman":
        return _roman_sequence(value)
    return str(value)


def _alpha_sequence(value: int, *, lowercase: bool) -> str:
    result: list[str] = []
    current = max(1, value)
    while current > 0:
        current -= 1
        result.append(chr(ord("a") + (current % 26)))
        current //= 26
    text = "".join(reversed(result))
    return text if lowercase else text.upper()


def _roman_sequence(value: int) -> str:
    numerals = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    current = max(1, value)
    parts: list[str] = []
    for number, numeral in numerals:
        while current >= number:
            parts.append(numeral)
            current -= number
    return "".join(parts)


def _xml_attr(element, name: str, default: str = "") -> str:
    if element is None:
        return default
    value = element.get(name)
    if value is None:
        return default
    return value


@dataclass(frozen=True, slots=True)
class _StyleState:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    all_caps: bool = False
    small_caps: bool = False
    superscript: bool = False
    subscript: bool = False


def _emit_paragraph_actions(actions: list[Action], paragraph: ParagraphBlock, current_style: _StyleState) -> _StyleState:
    if paragraph.prefix:
        current_style = _emit_style_transition(actions, current_style, _StyleState())
        actions.append(TypeText(paragraph.prefix))
    for run in paragraph.runs:
        current_style = _emit_style_transition(
            actions,
            current_style,
            _StyleState(
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                all_caps=run.all_caps,
                small_caps=run.small_caps,
                superscript=run.superscript,
                subscript=run.subscript,
            ),
        )
        if run.text:
            actions.append(TypeText(run.text))
    return current_style


def _emit_table_actions(actions: list[Action], table: TableBlock) -> _StyleState:
    current_style = _StyleState()
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row):
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                current_style = _emit_paragraph_actions(actions, paragraph, current_style)
                if paragraph_index + 1 < len(cell.paragraphs):
                    current_style = _emit_style_transition(actions, current_style, _StyleState())
                    actions.append(TypeText("\n"))
            if cell_index + 1 < len(row):
                current_style = _emit_style_transition(actions, current_style, _StyleState())
                actions.append(TypeText("\t"))
        if row_index + 1 < len(table.rows):
            current_style = _emit_style_transition(actions, current_style, _StyleState())
            actions.append(TypeText("\n"))
    current_style = _emit_style_transition(actions, current_style, _StyleState())
    return current_style


def _emit_style_transition(actions: list[Action], current: _StyleState, desired: _StyleState) -> _StyleState:
    if current == desired:
        return current

    for key, active_current, active_desired in (
        ("CTRL+EQUALS", current.subscript, desired.subscript),
        ("CTRL+SHIFT+EQUALS", current.superscript, desired.superscript),
        ("CTRL+SHIFT+K", current.small_caps, desired.small_caps),
        ("CTRL+SHIFT+A", current.all_caps, desired.all_caps),
        ("CTRL+U", current.underline, desired.underline),
        ("CTRL+I", current.italic, desired.italic),
        ("CTRL+B", current.bold, desired.bold),
    ):
        if active_current and not active_desired:
            actions.append(KeyPress(key))

    for key, active_current, active_desired in (
        ("CTRL+B", current.bold, desired.bold),
        ("CTRL+I", current.italic, desired.italic),
        ("CTRL+U", current.underline, desired.underline),
        ("CTRL+SHIFT+A", current.all_caps, desired.all_caps),
        ("CTRL+SHIFT+K", current.small_caps, desired.small_caps),
        ("CTRL+SHIFT+EQUALS", current.superscript, desired.superscript),
        ("CTRL+EQUALS", current.subscript, desired.subscript),
    ):
        if not active_current and active_desired:
            actions.append(KeyPress(key))

    return desired


def _style_close_actions(style: _StyleState) -> list[Action]:
    actions: list[Action] = []
    _emit_style_transition(actions, style, _StyleState())
    return actions


def _normalize_line_endings(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _run_has_formatting(run: TextRun) -> bool:
    return any(
        (
            run.bold,
            run.italic,
            run.underline,
            run.all_caps,
            run.small_caps,
            run.superscript,
            run.subscript,
        )
    )
